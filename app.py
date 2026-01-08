import streamlit as st
import pandas as pd
from pathlib import Path
import io
from docx import Document

# =========================================================
# 1. LOAD CONFIG FROM CSV FILES
# =========================================================
TIERS_PATH = Path("data/tiers.csv")
SHIPPING_PATH = Path("data/shipping.csv")
GLOBAL_PATH = Path("data/global_settings.csv")
PDLLA_TIERS_PATH = Path("data/Biogenomics Pricing - PDLLA.csv")


@st.cache_data
def load_config():
    tiers_df = pd.read_csv(TIERS_PATH)
    pdlla_df = pd.read_csv(PDLLA_TIERS_PATH)
    shipping_df = pd.read_csv(SHIPPING_PATH)
    global_df = pd.read_csv(GLOBAL_PATH)
    return tiers_df, pdlla_df, shipping_df, global_df


tiers_df, pdlla_df, shipping_df, global_df = load_config()

# global settings
GLOBAL_SETTINGS = {row["key"]: row["value"] for _, row in global_df.iterrows()}
CURRENCY = str(GLOBAL_SETTINGS.get("currency_symbol", "$"))
DEFAULT_MIN_CASES_GLOBAL = int(GLOBAL_SETTINGS.get("default_min_cases_global", 1))
DEFAULT_MAX_CASES_GLOBAL = int(GLOBAL_SETTINGS.get("default_max_cases_global", 500))

# build TIERS dict
TIERS_BASE = {}
for _, row in tiers_df.iterrows():
    TIERS_BASE[row["tier_name"]] = {
        "description": row["description"],
        "case_price": float(row["case_price"]),
        "cost_per_tx": float(row["cost_per_tx"]),
        "savings_vs_standard_pct": float(row["savings_vs_standard_pct"]),
        "tx_per_case": int(row["tx_per_case"]),
        "default_clinic_price_per_tx": float(row["default_clinic_price_per_tx"]),
        "default_extra_cost_per_tx": float(row["default_extra_cost_per_tx"]),
        "default_min_cases": int(row["default_min_cases"]),
        "default_max_cases": int(row["default_max_cases"]),
    }

TX_PER_CASE_DEFAULT = list(TIERS_BASE.values())[0]["tx_per_case"]

PDLLA_TIERS_BASE = {}
for _, row in pdlla_df.iterrows():
    PDLLA_TIERS_BASE[row["tier_name"]] = {
        "description": row["description"],
        "case_price": float(row["case_price"]),
        "cost_per_tx": float(row["cost_per_tx"]),
        "savings_vs_standard_pct": float(row["savings_vs_standard_pct"]),
        "tx_per_case": int(row["tx_per_case"]),
        "default_clinic_price_per_tx": float(row["default_clinic_price_per_tx"]),
        "default_extra_cost_per_tx": float(row["default_extra_cost_per_tx"]),
        "default_min_cases": int(row["default_min_cases"]),
        "default_max_cases": int(row["default_max_cases"]),
    }
    
# shipping dict
SHIPPING_BASE = {
    row["shipping_name"]: float(row["shipping_cost"])
    for _, row in shipping_df.iterrows()
}

# =========================================================
# 2. HELPERS
# =========================================================
def calc_roi(tier, num_cases, price_per_tx, extra_cost_per_tx, shipping_cost):
    case_price = tier["case_price"]
    cost_per_tx_product = tier["cost_per_tx"]
    tx_per_case = tier.get("tx_per_case", TX_PER_CASE_DEFAULT)

    total_cases = num_cases
    total_txs = total_cases * tx_per_case

    product_cost = total_cases * case_price
    total_cost = product_cost + shipping_cost

    total_cost_per_tx = cost_per_tx_product + extra_cost_per_tx
    revenue_per_tx = price_per_tx
    profit_per_tx = revenue_per_tx - total_cost_per_tx

    total_revenue = revenue_per_tx * total_txs
    total_profit = profit_per_tx * total_txs

    margin_pct = (total_profit / total_revenue * 100) if total_revenue else 0
    roi_pct = (total_profit / total_cost * 100) if total_cost else 0

    breakeven_txs = total_cost / profit_per_tx if profit_per_tx > 0 else None

    return {
        "total_cases": total_cases,
        "total_txs": total_txs,
        "product_cost": product_cost,
        "shipping_cost": shipping_cost,
        "total_cost": total_cost,
        "revenue_per_tx": revenue_per_tx,
        "cost_per_tx_product": cost_per_tx_product,
        "extra_cost_per_tx": extra_cost_per_tx,
        "total_cost_per_tx": total_cost_per_tx,
        "profit_per_tx": profit_per_tx,
        "total_revenue": total_revenue,
        "total_profit": total_profit,
        "margin_pct": margin_pct,
        "roi_pct": roi_pct,
        "breakeven_txs": breakeven_txs,
    }


def fc(x):
    return f"{CURRENCY}{x:,.0f}"


def fc1(x):
    return f"{CURRENCY}{x:,.1f}"


def build_word_report(
    genovia_tier_name,
    genovia_results,
    pdlla_tier_name,
    pdlla_results,
    num_cases,
    price_per_tx,
    extra_cost_per_tx,
    shipping_name,
    shipping_cost,
    genovia_comp_df,
):
    """Create a Word document (as BytesIO) summarizing the scenario + comparison."""
    doc = Document()

    # Title
    doc.add_heading("Genovia vs PDLLA ROI Report", level=1)
    doc.add_paragraph(
        "This report summarizes the ROI scenario generated from the Genovia ROI Calculator, "
        "comparing Genovia exosome tiers with PDLLA tiers under the same clinic assumptions."
    )

    # Section 1 – Scenario Overview
    doc.add_heading("1. Scenario Overview", level=2)
    doc.add_paragraph(f"Number of cases: {num_cases}")
    doc.add_paragraph(f"Clinic price per treatment: {fc1(price_per_tx)}")
    doc.add_paragraph(f"Other cost per treatment: {fc1(extra_cost_per_tx)}")
    doc.add_paragraph(f"Shipping option: {shipping_name} ({fc(shipping_cost)})")
    doc.add_paragraph(f"Genovia tier: {genovia_tier_name}")
    doc.add_paragraph(f"PDLLA tier: {pdlla_tier_name}")


    # Section 2 – Key Financial Metrics
    doc.add_heading("2. Genovia Financial Metrics", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    metrics = [
        ("Total treatments", f"{genovia_results['total_txs']:,}"),
        ("Total revenue", fc(genovia_results["total_revenue"])),
        ("Total cost (product + shipping)", fc(genovia_results["total_cost"])),
        ("Total profit", fc(genovia_results["total_profit"])),
        ("Profit per treatment", fc1(genovia_results["profit_per_tx"])),
        ("Profit margin", f"{genovia_results['margin_pct']:.1f}%"),
        ("ROI on order", f"{genovia_results['roi_pct']:.1f}%"),
    ]
    for name, value in metrics:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = value

    # Section 3 – Interpretation
    doc.add_heading("3. PDLLA Financial Metrics", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Metric"
    hdr2[1].text = "Value"
    metrics_other = [
        ("Total treatments", f"{pdlla_results['total_txs']:,}"),
        ("Total revenue", fc(pdlla_results["total_revenue"])),
        ("Total cost (product + shipping)", fc(pdlla_results["total_cost"])),
        ("Total profit", fc(pdlla_results["total_profit"])),
        ("Profit per treatment", fc1(pdlla_results["profit_per_tx"])),
        ("Profit margin", f"{pdlla_results['margin_pct']:.1f}%"),
        ("ROI on order", f"{pdlla_results['roi_pct']:.1f}%"),
    ]
    for name, value in metrics_other:
        row = table2.add_row().cells
        row[0].text = name
        row[1].text = value

    # Section 4 – Tier Comparison
    doc.add_heading("4. Genovia vs PDLLA – Direct Comparison", level=2)
    delta_profit = genovia_results["total_profit"] - pdlla_results["total_profit"]
    delta_roi = genovia_results["roi_pct"] - pdlla_results["roi_pct"]

    doc.add_paragraph(
        f"Under the same clinic assumptions, Genovia generates {fc(genovia_results['total_profit'])} in total profit "
        f"versus {fc(pdlla_results['total_profit'])} for PDLLA."
    )
    doc.add_paragraph(
        f"Profit difference (Genovia − PDLLA): {fc(delta_profit)}."
    )
    doc.add_paragraph(
        f"ROI difference (Genovia − PDLLA): {delta_roi:.1f} percentage points."
    )

    # Section 5 – Recommendations
    doc.add_heading("5. Genovia Tier Comparison at Same Clinic Price", level=2)
    comp_table = doc.add_table(rows=1, cols=4)
    ch = comp_table.rows[0].cells
    ch[0].text = "Tier"
    ch[1].text = "Cost per Treatment (Genovia)"
    ch[2].text = "Total Profit"
    ch[3].text = "ROI %"

    for _, row in genovia_comp_df.iterrows():
        r = comp_table.add_row().cells
        r[0].text = str(row["Tier"])
        r[1].text = fc1(row["Cost per treatment"])
        r[2].text = fc(row["Total Profit"])
        r[3].text = f"{row['ROI %']:.1f}%"

    # 6. Recommendations
    doc.add_heading("6. Recommendations", level=2)
    recs = [
        "Use this model live in sales conversations to show scenario-based ROI for both Genovia and PDLLA.",
        "Highlight Genovia tiers that deliver the highest ROI at the clinic’s expected pricing and volume.",
        "Provide customized ROI reports to priority accounts to support tier and product selection.",
    ]
    for r in recs:
        doc.add_paragraph(f"• {r}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



# =========================================================
# 3. UI LAYOUT
# =========================================================
st.set_page_config(page_title="Genovia ROI Calculator", page_icon="💧", layout="centered")

st.title("Genovia™ ROI Calculator")
st.caption("Compare Genovia exosome tiers and PDLLA tiers under the same clinic assumptions.")

# ------------------ SIDEBAR: ALL INPUTS ------------------
with st.sidebar:
    st.markdown("### Tier Settings & Pricing Inputs (Genovia)")
    st.caption(
        "Genovia tier pricing loads from tiers.csv. Adjust below for scenario modeling. "
        "Changes affect only this session."
    )


    tiers_runtime = {k: v.copy() for k, v in TIERS_BASE.items()}
    shipping_runtime = SHIPPING_BASE.copy()

    # Tier settings
    for tier_name, tier in tiers_runtime.items():
        with st.expander(f"{tier_name} — Genovia Pricing Settings", expanded=False):
            tier["case_price"] = st.number_input(
                f"{tier_name} · Case Price",
                min_value=0.0,
                step=10.0,
                value=tier["case_price"],
                key=f"case_price_{tier_name}",
            )
            tier["cost_per_tx"] = st.number_input(
                f"{tier_name} · Cost per Treatment (Genovia)",
                min_value=0.0,
                step=1.0,
                value=tier["cost_per_tx"],
                key=f"cost_per_tx_{tier_name}",
            )
            tier["tx_per_case"] = st.number_input(
                f"{tier_name} · Treatments per Case",
                min_value=1,
                step=1,
                value=tier["tx_per_case"],
                key=f"tx_per_case_{tier_name}",
            )

    st.markdown("---")
    st.markdown("### Clinic Inputs")

    tier_choice = st.selectbox("Genovia tier offered", list(tiers_runtime.keys()))
    tier_selected = tiers_runtime[tier_choice]

    num_cases = st.number_input(
        "Number of cases in this order",
        min_value=int(tier_selected["default_min_cases"]),
        max_value=int(tier_selected["default_max_cases"]),
        value=int(tier_selected["default_min_cases"]),
        step=1,
    )

    price_per_tx = st.number_input(
        "Clinic price per treatment ($)",
        value=float(tier_selected["default_clinic_price_per_tx"]),
        min_value=0.0,
        step=50.0,
    )

    extra_cost_per_tx = st.number_input(
        "Other per-treatment cost (tips, etc.)",
        value=float(tier_selected["default_extra_cost_per_tx"]),
        min_value=0.0,
        step=10.0,
    )

    st.markdown("### Shipping Assumptions")
    shipping_name = st.selectbox("Shipping option", list(shipping_runtime.keys()))
    # allow adjusting shipping cost in sidebar as well
    shipping_cost = st.number_input(
        "Shipping cost for this order",
        min_value=0.0,
        step=5.0,
        value=shipping_runtime[shipping_name],
        key=f"shipping_cost_active",
    )
    
    st.markdown("---")
    st.markdown("### PDLLA Tier for Comparison")
    pdlla_tier_choice = st.selectbox("PDLLA tier", list(PDLLA_TIERS_BASE.keys()))
    pdlla_tier_selected = PDLLA_TIERS_BASE[pdlla_tier_choice]



# ------------------ MAIN: ROI SUMMARY & OUTPUTS ------------------
genovia_results = calc_roi(
    tier=tier_selected,
    num_cases=int(num_cases),
    price_per_tx=price_per_tx,
    extra_cost_per_tx=extra_cost_per_tx,
    shipping_cost=shipping_cost,
)

# --- PDLLA results (same clinic assumptions) ---
pdlla_results = calc_roi(
    tier=pdlla_tier_selected,
    num_cases=int(num_cases),
    price_per_tx=price_per_tx,
    extra_cost_per_tx=extra_cost_per_tx,
    shipping_cost=shipping_cost,
)

st.subheader("Step 2 — ROI Summary: Genovia vs PDLLA")

left, right = st.columns(2)

with left:
    st.markdown(f"#### Genovia – {tier_choice} Tier")
    col1, col2, col3 = st.columns(3)
    col1.metric("Total Revenue", fc(genovia_results["total_revenue"]))
    col2.metric("Total Cost", fc(genovia_results["total_cost"]))
    col3.metric("Total Profit", fc(genovia_results["total_profit"]))

    col4, col5 = st.columns(2)
    col4.metric("Margin", f"{genovia_results['margin_pct']:.1f}%")
    col5.metric("ROI", f"{genovia_results['roi_pct']:.1f}%")

with right:
    st.markdown(f"#### PDLLA – {pdlla_tier_choice} Tier")
    c1, c2, c3 = st.columns(3)
    c1.metric("Total Revenue", fc(pdlla_results["total_revenue"]))
    c2.metric("Total Cost", fc(pdlla_results["total_cost"]))
    c3.metric("Total Profit", fc(pdlla_results["total_profit"]))

    c4, c5 = st.columns(2)
    c4.metric("Margin", f"{pdlla_results['margin_pct']:.1f}%")
    c5.metric("ROI", f"{pdlla_results['roi_pct']:.1f}%")

summary_df = pd.DataFrame(
    {
        "Metric": ["Total Revenue", "Total Cost", "Total Profit"],
        "Genovia": [
            genovia_results["total_revenue"],
            genovia_results["total_cost"],
            genovia_results["total_profit"],
        ],
        "PDLLA": [
            pdlla_results["total_revenue"],
            pdlla_results["total_cost"],
            pdlla_results["total_profit"],
        ],
    }
).set_index("Metric")
st.bar_chart(summary_df)

scenario_export_df = pd.DataFrame(
    [
        {
            "Genovia tier": tier_choice,
            "PDLLA tier": pdlla_tier_choice,
            "Number of cases": num_cases,
            "Clinic price per treatment": price_per_tx,
            "Other cost per treatment": extra_cost_per_tx,
            "Shipping option": shipping_name,
            "Shipping cost": shipping_cost,
            "Genovia total revenue": genovia_results["total_revenue"],
            "Genovia total cost": genovia_results["total_cost"],
            "Genovia total profit": genovia_results["total_profit"],
            "Genovia margin %": genovia_results["margin_pct"],
            "Genovia ROI %": genovia_results["roi_pct"],
            "PDLLA total revenue": pdlla_results["total_revenue"],
            "PDLLA total cost": pdlla_results["total_cost"],
            "PDLLA total profit": pdlla_results["total_profit"],
            "PDLLA margin %": pdlla_results["margin_pct"],
            "PDLLA ROI %": pdlla_results["roi_pct"],
        }
    ]
)


st.download_button(
    label="⬇️ Download current scenario (CSV)",
    data=scenario_export_df.to_csv(index=False),
    file_name="genovia_pdlla_roi_scenario.csv",
    mime="text/csv",
)

st.markdown("---")

st.subheader("Genovia – Compare Tiers at Same Clinic Price")

comparison = []
for tname, info in tiers_runtime.items():
    r = calc_roi(
        tier=info,
        num_cases=int(num_cases),
        price_per_tx=price_per_tx,
        extra_cost_per_tx=extra_cost_per_tx,
        shipping_cost=shipping_cost,
    )
    comparison.append(
        {
            "Tier": tname,
            "Cost per treatment": r["cost_per_tx_product"],
            "Total Profit": r["total_profit"],
            "ROI %": r["roi_pct"],
        }
    )

genovia_comp_df = pd.DataFrame(comparison)
st.table(genovia_comp_df)

st.markdown("#### Genovia – Profit by Tier")
st.bar_chart(genovia_comp_df.set_index("Tier")["Total Profit"])

st.markdown("#### Genovia – ROI by Tier")
st.bar_chart(genovia_comp_df.set_index("Tier")["ROI %"])

st.download_button(
    label="⬇️ Download Genovia tier comparison (CSV)",
    data=genovia_comp_df.to_csv(index=False),
    file_name="genovia_tier_comparison.csv",
    mime="text/csv",
)

report_buffer = build_word_report(
    genovia_tier_name=tier_choice,
    genovia_results=genovia_results,
    pdlla_tier_name=pdlla_tier_choice,
    pdlla_results=pdlla_results,
    num_cases=int(num_cases),
    price_per_tx=price_per_tx,
    extra_cost_per_tx=extra_cost_per_tx,
    shipping_name=shipping_name,
    shipping_cost=shipping_cost,
    genovia_comp_df=genovia_comp_df,
)

st.markdown("---")
st.download_button(
    label="⬇️ Download Word report (Genovia vs PDLLA)",
    data=report_buffer,
    file_name="Genovia_vs_PDLLA_ROI_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
